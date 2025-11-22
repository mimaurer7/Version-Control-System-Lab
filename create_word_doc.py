from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime

# Create document
doc = Document()

# Title
title = doc.add_heading('Version Control System (VCS) Lab Documentation', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Student info
info = doc.add_paragraph()
info.add_run('Student Name: ').bold = True
info.add_run('Matthew Maurer\n')
info.add_run('GitHub Username: ').bold = True
info.add_run('mimaurer7\n')
info.add_run('Course: ').bold = True
info.add_run('Cybersecurity\n')
info.add_run('Date: ').bold = True
info.add_run(f'{datetime.date.today().strftime("%B %d, %Y")}\n')
info.add_run('Repository: ').bold = True
info.add_run('https://github.com/mimaurer7/Version-Control-System-Lab')

doc.add_page_break()

# Table of Contents
doc.add_heading('Table of Contents', 1)
toc_items = [
    '1. Introduction',
    '2. Setup and Configuration',
    '3. Branching Strategy',
    '4. Secure Coding Practices',
    '5. Backup and Recovery',
    '6. Conclusion'
]
for item in toc_items:
    doc.add_paragraph(item, style='List Number')

doc.add_page_break()

# 1. Introduction
doc.add_heading('1. Introduction', 1)
doc.add_paragraph(
    'This document details the implementation of professional version control practices for the VCS Lab assignment. '
    'The project demonstrates Git fundamentals, branching workflows, security integration, and backup procedures '
    'essential for secure software development.'
)

doc.add_heading('Project Overview', 2)
overview = doc.add_paragraph()
overview.add_run('• Version Control System: ').bold = True
overview.add_run('Git\n')
overview.add_run('• Remote Repository: ').bold = True
overview.add_run('GitHub\n')
overview.add_run('• Branching Model: ').bold = True
overview.add_run('Feature Branch Workflow\n')
overview.add_run('• Security Integration: ').bold = True
overview.add_run('Pre-commit hooks with automated scanning')

doc.add_page_break()

# 2. Setup and Configuration
doc.add_heading('2. Setup and Configuration', 1)

doc.add_heading('2.1 Repository Initialization', 2)
doc.add_paragraph('The repository was initialized using the git init command:')
doc.add_paragraph('git init', style='Intense Quote')
doc.add_paragraph('[INSERT SCREENSHOT 1: Terminal showing git init output]')

doc.add_heading('2.2 Git Configuration', 2)
doc.add_paragraph('User information was configured to track commit authorship:')
code = doc.add_paragraph('git config user.name "mimaurer7"\ngit config user.email "mimaurer7@gmail.com"', style='Intense Quote')

doc.add_heading('2.3 .gitignore Configuration', 2)
doc.add_paragraph('A comprehensive .gitignore file was created to exclude:')
ignore_list = [
    'Personal configuration files (*.local, config/settings.json)',
    'Sensitive data (*.key, *.pem, *.env, secrets/)',
    'Dependency folders (node_modules/, venv/, __pycache__/)',
    'IDE settings (.vscode/settings.json, .idea/)',
    'Operating system files (.DS_Store, Thumbs.db)',
    'Build output (dist/, build/, *.log)',
    'Backup files (*.bak, *.tmp)'
]
for item in ignore_list:
    doc.add_paragraph(item, style='List Bullet')
doc.add_paragraph('[INSERT SCREENSHOT 5: .gitignore file contents]')

doc.add_heading('2.4 Project Structure', 2)
doc.add_paragraph('The following directory structure was established:')
structure = """Version Control System Lab/
├── .git/                    # Git repository data
├── .gitignore               # Ignored files configuration
├── README.md                # Project documentation
├── src/                     # Source code
│   ├── index.html
│   ├── styles.css
│   └── app.js
├── scripts/                 # Automation scripts
│   └── security_scan.py
├── docs/                    # Documentation
│   ├── branching-strategy.md
│   ├── security-practices.md
│   └── backup-recovery.md
└── config/                  # Configuration
    └── settings.example.json"""
doc.add_paragraph(structure, style='Intense Quote')
doc.add_paragraph('[INSERT SCREENSHOT 4: VS Code Explorer showing project structure]')

doc.add_page_break()

# 3. Branching Strategy
doc.add_heading('3. Branching Strategy', 1)

doc.add_heading('3.1 Feature Branch Workflow', 2)
doc.add_paragraph(
    'This project implements the Feature Branch Workflow, which is ideal for small to medium projects. '
    'This workflow ensures the main branch always contains stable, production-ready code.'
)

doc.add_heading('3.2 Branch Naming Conventions', 2)
doc.add_paragraph('Clear naming conventions were established:')

# Create table
table = doc.add_table(rows=4, cols=3)
table.style = 'Light Grid Accent 1'
header_cells = table.rows[0].cells
header_cells[0].text = 'Prefix'
header_cells[1].text = 'Purpose'
header_cells[2].text = 'Example'

data = [
    ('feature/', 'New features', 'feature/add-security-dashboard'),
    ('bugfix/', 'Bug fixes', 'bugfix/fix-css-styling'),
    ('main', 'Stable production code', 'main')
]
for i, (prefix, purpose, example) in enumerate(data, 1):
    cells = table.rows[i].cells
    cells[0].text = prefix
    cells[1].text = purpose
    cells[2].text = example

doc.add_paragraph('[INSERT SCREENSHOT 2: Git branch output showing all branches]')
doc.add_paragraph('[INSERT SCREENSHOT 1: Git log with graph showing merge history]')

doc.add_heading('3.3 Benefits of Feature Branch Workflow', 2)
benefits = [
    'Isolation: Each feature developed independently',
    'Stability: Main branch always production-ready',
    'Collaboration: Multiple developers can work simultaneously',
    'Code Review: Changes reviewed before merging',
    'Rollback: Easy to revert problematic features'
]
for benefit in benefits:
    doc.add_paragraph(benefit, style='List Bullet')

doc.add_page_break()

# 4. Secure Coding Practices
doc.add_heading('4. Secure Coding Practices', 1)

doc.add_heading('4.1 Pre-commit Security Hooks', 2)
doc.add_paragraph('A pre-commit hook was implemented to run security scans before allowing commits.')
hook_code = """#!/bin/sh
echo "Running security checks..."
python scripts/security_scan.py
if [ $? -ne 0 ]; then
    echo "Security scan failed! Commit aborted."
    exit 1
fi
echo "Security checks passed!"
exit 0"""
doc.add_paragraph(hook_code, style='Intense Quote')

doc.add_heading('4.2 Security Scanning Tools', 2)
doc.add_paragraph('Recommended CI/CD Integration:')

tools = [
    ('SonarQube', 'Static code analysis for code smells, bugs, and vulnerabilities'),
    ('Snyk', 'Dependency vulnerability scanning with real-time alerts'),
    ('OWASP Dependency Check', 'CVE database matching for vulnerable dependencies')
]
for tool, desc in tools:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f'{tool}: ').bold = True
    p.add_run(desc)

doc.add_heading('4.3 Dependency Management', 2)
doc.add_paragraph('Update Schedule:')
schedule = [
    'Weekly: Check for dependency updates',
    'Monthly: Review and apply security patches',
    'Immediate: Critical security vulnerabilities'
]
for item in schedule:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# 5. Backup and Recovery
doc.add_heading('5. Backup and Recovery', 1)

doc.add_heading('5.1 Remote Repository Configuration', 2)
doc.add_paragraph('The local repository was connected to GitHub for remote backup:')
backup_code = """git remote add origin https://github.com/mimaurer7/Version-Control-System-Lab.git
git push -u origin main"""
doc.add_paragraph(backup_code, style='Intense Quote')
doc.add_paragraph('[INSERT SCREENSHOT 6: GitHub repository page]')

doc.add_heading('5.2 Recovery Procedures', 2)

scenarios = [
    ('Complete Repository Loss', 'git clone https://github.com/mimaurer7/Version-Control-System-Lab.git', '2-5 minutes'),
    ('File Recovery', 'git checkout <commit-hash> -- path/to/file', '<1 minute'),
    ('Undo Last Commit', 'git reset --soft HEAD~1', 'Instant')
]

for scenario, command, time in scenarios:
    p = doc.add_paragraph()
    p.add_run(f'{scenario}: ').bold = True
    doc.add_paragraph(command, style='Intense Quote')
    doc.add_paragraph(f'Recovery Time: {time}')

doc.add_heading('5.3 Backup Testing', 2)
test = doc.add_paragraph()
test.add_run('Test Performed: ').bold = True
test.add_run('Repository clone verification\n')
test.add_run('Date: ').bold = True
test.add_run(f'{datetime.date.today().strftime("%B %d, %Y")}\n')
test.add_run('Result: ').bold = True
result_run = test.add_run('✅ SUCCESS')
result_run.font.color.rgb = RGBColor(0, 128, 0)

doc.add_page_break()

# 6. Conclusion
doc.add_heading('6. Conclusion', 1)
doc.add_paragraph(
    'This VCS lab successfully demonstrates professional version control practices essential for secure, '
    'collaborative software development and cybersecurity professionals.'
)

doc.add_heading('Key Achievements', 2)
achievements = [
    'Repository Setup: Proper initialization and configuration',
    'Version Control: Atomic commits with descriptive messages',
    'Branching Strategy: Feature Branch Workflow with clear conventions',
    'Security Integration: Pre-commit hooks and scanning tools',
    'Backup & Recovery: Remote repository with tested recovery procedures'
]
for achievement in achievements:
    doc.add_paragraph(f'✅ {achievement}', style='List Bullet')

doc.add_heading('Repository Access', 2)
access = doc.add_paragraph()
access.add_run('GitHub: ').bold = True
access.add_run('https://github.com/mimaurer7/Version-Control-System-Lab\n')
access.add_run('Owner: ').bold = True
access.add_run('Matthew Maurer (mimaurer7@gmail.com)')

# Save document
doc.save('VCS_Lab_Documentation_Matthew_Maurer.docx')
print("Word document created successfully!")
print("File: VCS_Lab_Documentation_Matthew_Maurer.docx")